import json
import os
import subprocess
import sys
from datetime import datetime
from io import BytesIO
from pathlib import Path

from flask import Flask, jsonify, redirect, render_template_string, request, send_file, url_for
from supabase import create_client
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from openpyxl import Workbook

app = Flask(__name__)
BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = BASE_DIR / "cover_letter_template.docx"
LOG_FILE = BASE_DIR / "run.log"
META_FILE = BASE_DIR / "job_meta.json"
RUN_STATE_FILE = BASE_DIR / "run_state.json"
SUPABASE_TABLE = "jobs"
STATUS_ORDER = ["not_applied", "applied", "not_to_apply"]
STATUS_LABELS = {
    "not_applied": "Nicht beworben",
    "applied": "Beworben",
    "not_to_apply": "Nicht bewerben",
}
SEARCH_TERMS_FILE = BASE_DIR / "search_terms.json"
LETTER_FONT_NAME = "Helvetica"
LETTER_FONT_SIZE = 9
LETTER_SPACE_AFTER_PT = 6
LETTER_TEXT_COLOR = RGBColor(29, 39, 49)
LETTER_LINE_SPACING = 1.5

current_process = None


def is_running():
    global current_process
    refresh_run_state()
    return current_process is not None and current_process.poll() is None


def load_json(path, default):
    if not path.exists():
        return default
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return default


def save_json(path, data):
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def default_search_terms():
    return [
        "Türkisch",
        "Migration",
        "Geflüchtete",
        "Soziologie",
        "Integration",
        "Farsi",
        "Sozialbetreuer",
        "Persisch",
        "Soziologie",
        "Internationale Beziehungen",
        "Sozialwissenschaftler",
        "Flüchtlinge",
        "Mehrsprachigkeit",
        "Zuwanderer",
    ]


def repair_mojibake(text):
    raw = str(text or "").strip()
    if not raw:
        return ""
    try:
        repaired = raw.encode("latin-1").decode("utf-8")
        return repaired
    except (UnicodeEncodeError, UnicodeDecodeError):
        return raw


def load_search_terms():
    data = load_json(SEARCH_TERMS_FILE, {"terms": default_search_terms()})
    raw_terms = data.get("terms", data) if isinstance(data, dict) else data
    terms = [repair_mojibake(term) for term in raw_terms or []]
    terms = [term for term in terms if term]
    return terms or default_search_terms()


def save_search_terms_text(text):
    terms = [term.strip() for term in text.replace(";", "\n").replace(",", "\n").splitlines()]
    terms = [repair_mojibake(term) for term in terms]
    terms = [term for term in terms if term]
    save_json(SEARCH_TERMS_FILE, {"terms": terms})
    return terms


def load_run_state():
    return load_json(
        RUN_STATE_FILE,
        {
            "status": "idle",
            "started_at": "",
            "finished_at": "",
            "returncode": None,
        },
    )


def save_run_state(status, *, returncode=None):
    state = load_run_state()
    state["status"] = status
    if status == "running":
        state["started_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        state["finished_at"] = ""
        state["returncode"] = None
    else:
        state["finished_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        state["returncode"] = returncode
    save_json(RUN_STATE_FILE, state)


def refresh_run_state():
    global current_process
    if current_process is None:
        return
    returncode = current_process.poll()
    if returncode is None:
        return
    save_run_state("completed" if returncode == 0 else "failed", returncode=returncode)
    current_process = None


def clean_log_content(content):
    noisy_fragments = (
        "Serving Flask app",
        "Debug mode:",
        "WARNING: This is a development server",
        "Running on http://",
        "Press CTRL+C to quit",
        "Restarting with stat",
        "Debugger is active",
        "Debugger PIN:",
        '"HEAD / HTTP/1.1"',
        '"GET / HTTP/1.1"',
    )
    cleaned_lines = []
    for line in content.splitlines():
        if any(fragment in line for fragment in noisy_fragments):
            continue
        cleaned_lines.append(line)
    cleaned = "\n".join(cleaned_lines).strip()
    return cleaned or "No job-search logs yet. Run the search first."


def get_supabase_client():
    url = os.environ.get("SUPABASE_URL")
    key = os.environ.get("SUPABASE_ANON_KEY")
    if not url or not key:
        return None
    return create_client(url, key)


def normalize_path_string(path_string):
    if not path_string:
        return None
    relative = Path(str(path_string).replace("\\", "/"))
    return BASE_DIR / relative


def rel_download_path(path):
    return str(path.relative_to(BASE_DIR)).replace("\\", "/")


def date_only(value):
    text = str(value or "").strip()
    if not text:
        return ""
    return text[:10]


def load_jobs():
    client = get_supabase_client()
    if client is None:
        return []
    rows = []
    page_size = 1000
    offset = 0
    while True:
        response = (
            client.table(SUPABASE_TABLE)
            .select("*")
            .order("created_at", desc=True)
            .range(offset, offset + page_size - 1)
            .execute()
        )
        batch = response.data or []
        rows.extend(batch)
        if len(batch) < page_size:
            break
        offset += page_size

    jobs = []
    for row in rows:
        refnr = str(row.get("refnr", "")).strip()
        if not refnr:
            continue
        cover_letter_path = normalize_path_string(row.get("cover_letter_path"))
        job_description_path = normalize_path_string(row.get("job_description_path"))
        application_status = row.get("application_status", "not_applied")
        if application_status == "to_apply":
            application_status = "not_applied"
        jobs.append(
            {
                "refnr": refnr,
                "job_id": row.get("job_id") or "",
                "date": row.get("date") or "",
                "keyword_score": row.get("keyword_score") or "",
                "ai_match_score": row.get("ai_match_score") or "",
                "title": row.get("title") or "",
                "employer": row.get("employer") or "",
                "employer_street": row.get("employer_street") or "",
                "employer_postal_code": row.get("employer_postal_code") or "",
                "employer_city": row.get("employer_city") or "",
                "city": row.get("city") or "",
                "reason": row.get("reason") or "",
                "job_url": row.get("job_url") or "",
                "cover_letter_path": cover_letter_path,
                "job_description_path": job_description_path,
                "cover_letter_text": row.get("cover_letter_text") or "",
                "job_description_text": row.get("job_description_text") or "",
                "has_cover_letter": bool(row.get("has_cover_letter")),
                "application_status": application_status,
                "application_status_label": STATUS_LABELS.get(
                    application_status,
                    "Not Applied",
                ),
                "application_method": row.get("application_method", ""),
                "application_result": row.get("application_result", ""),
                "applied_at": row.get("applied_at", ""),
                "note": row.get("note", ""),
                "created_at": row.get("created_at", ""),
                "updated_at": row.get("updated_at", ""),
            }
        )

    def sort_key(item):
        return (
            str(item["created_at"] or ""),
            str(item["refnr"] or ""),
        )

    jobs.sort(key=sort_key, reverse=True)
    return jobs


def fetch_job(refnr):
    client = get_supabase_client()
    if client is None:
        return None
    response = client.table(SUPABASE_TABLE).select("*").eq("refnr", str(refnr)).limit(1).execute()
    rows = response.data or []
    return rows[0] if rows else None


def fetch_job_by_job_id(job_id):
    client = get_supabase_client()
    if client is None:
        return None
    response = client.table(SUPABASE_TABLE).select("*").eq("job_id", str(job_id)).limit(1).execute()
    rows = response.data or []
    return rows[0] if rows else None


def german_date_string():
    months = [
        "Januar",
        "Februar",
        "Maerz",
        "April",
        "Mai",
        "Juni",
        "Juli",
        "August",
        "September",
        "Oktober",
        "November",
        "Dezember",
    ]
    today = datetime.now()
    return f"{today.day}. {months[today.month - 1]} {today.year}"


def style_paragraph(paragraph, *, space_after_pt=LETTER_SPACE_AFTER_PT):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(space_after_pt)
    paragraph.paragraph_format.line_spacing = LETTER_LINE_SPACING
    for run in paragraph.runs:
        run.font.name = LETTER_FONT_NAME
        run.font.size = Pt(LETTER_FONT_SIZE)
        run.font.color.rgb = LETTER_TEXT_COLOR


def style_table_paragraph(paragraph):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = LETTER_LINE_SPACING
    for run in paragraph.runs:
        run.font.name = LETTER_FONT_NAME
        run.font.size = Pt(LETTER_FONT_SIZE)
        run.font.color.rgb = LETTER_TEXT_COLOR


def replace_paragraph_text(paragraph, text, *, bold=None, space_after_pt=LETTER_SPACE_AFTER_PT):
    paragraph.text = ""
    run = paragraph.add_run(text)
    if bold is not None:
        run.bold = bold
    style_paragraph(paragraph, space_after_pt=space_after_pt)


def insert_paragraph_after(paragraph, text):
    new_paragraph = paragraph.insert_paragraph_before("")
    paragraph._p.addnext(new_paragraph._p)
    new_paragraph.style = paragraph.style
    if paragraph.alignment is not None:
        new_paragraph.alignment = paragraph.alignment
    new_paragraph.add_run(text)
    style_paragraph(new_paragraph)
    return new_paragraph


def restyle_document(document):
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            style_paragraph(paragraph)

    for table in document.tables:
        for row_cells in table.rows:
            row_cells.height = None
            row_cells.height_rule = None
            for cell in row_cells.cells:
                for paragraph in cell.paragraphs:
                    style_table_paragraph(paragraph)


def fill_cover_letter_template(document, row):
    subject = f"Bewerbung als {row.get('title') or 'Stelle'}"
    body_text = (row.get("cover_letter_text") or "").strip()
    body_parts = [part.strip() for part in body_text.split("\n\n") if part.strip()] or [body_text]

    employer = row.get("employer") or ""
    city = row.get("city") or ""
    street = row.get("employer_street") or ""
    postal_city = " ".join(
        part
        for part in [row.get("employer_postal_code"), row.get("employer_city") or city]
        if part
    ).strip()

    replacements = {
        "{{DATE}}": german_date_string(),
        "{{SUBJECT}}": subject,
        "{{EMPLOYER}}": employer,
        "{{STREET, HOUSE NUMBER}}": street,
        "{{POSTAL NUMBER, CITY}}": postal_city,
    }

    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text in replacements:
            replace_paragraph_text(
                paragraph,
                replacements[text],
                bold=(text == "{{SUBJECT}}"),
            )
        elif text == "{{BODY}}":
            replace_paragraph_text(paragraph, body_parts[0] if body_parts else "")
            current_paragraph = paragraph
            for extra in body_parts[1:]:
                current_paragraph = insert_paragraph_after(current_paragraph, extra)

    for table in document.tables:
        for row_cells in table.rows:
            row_cells.height = None
            row_cells.height_rule = None
            for cell in row_cells.cells:
                for paragraph in cell.paragraphs:
                    text = (paragraph.text or "").strip()
                    if text in replacements:
                        replace_paragraph_text(paragraph, replacements[text], space_after_pt=0)
                        style_table_paragraph(paragraph)
                    else:
                        style_table_paragraph(paragraph)

    restyle_document(document)


def group_jobs_by_status(jobs):
    grouped = {status: [] for status in STATUS_ORDER}
    for job in jobs:
        status = job.get("application_status") or "not_applied"
        if status == "to_apply":
            status = "not_applied"
        if status not in grouped:
            status = "not_applied"
        grouped[status].append(job)

    return [
        {
            "key": status,
            "label": STATUS_LABELS[status],
            "jobs": grouped[status],
        }
        for status in STATUS_ORDER
    ]


@app.get("/")
def dashboard():
    return redirect(url_for("search_page"))

    # Legacy dashboard kept below but bypassed by redirect.
    jobs = load_jobs()
    grouped_sections = group_jobs_by_status(jobs)
    application_results = [job for job in jobs if (job.get("application_result") or "").strip()]
    running = is_running()
    run_state = load_run_state()
    status = "Running" if running else run_state.get("status", "idle").title()
    logs = LOG_FILE.read_text(encoding="utf-8", errors="replace") if LOG_FILE.exists() else "No log file yet."
    recent_logs = "\n".join(logs.splitlines()[-40:])

    template = """
    <!doctype html>
    <html lang="en">
    <head>
      <meta charset="utf-8">
      <meta name="viewport" content="width=device-width, initial-scale=1">
      <title>Job Automation</title>
      {% if status == "Running" %}
      <meta http-equiv="refresh" content="5">
      {% endif %}
      <style>
        :root {
          --bg: #eef3f8;
          --panel: #ffffff;
          --ink: #18324b;
          --muted: #5b7288;
          --line: #d7e1ea;
          --accent: #0f5ea8;
          --accent-2: #dbeafe;
          --success: #1f7a4d;
          --warn: #9a6700;
          --shadow: 0 18px 50px rgba(17, 36, 56, 0.10);
          --radius: 18px;
        }
        * { box-sizing: border-box; }
        body {
          margin: 0;
          font-family: "Segoe UI", Arial, sans-serif;
          background:
            radial-gradient(circle at top left, #f8fbff 0, #eef3f8 48%, #e8eef5 100%);
          color: var(--ink);
        }
        .wrap {
          max-width: 1180px;
          margin: 0 auto;
          padding: 20px 14px 40px;
        }
        .hero {
          background: linear-gradient(135deg, #113350, #245f95);
          color: white;
          border-radius: 24px;
          padding: 22px 20px;
          box-shadow: var(--shadow);
          margin-bottom: 18px;
        }
        .hero h1 {
          margin: 0 0 10px;
          font-size: 28px;
        }
        .hero p {
          margin: 0;
          color: rgba(255,255,255,0.86);
        }
        .toolbar {
          display: grid;
          grid-template-columns: repeat(auto-fit, minmax(220px, 1fr));
          gap: 14px;
          margin: 18px 0;
        }
        .panel {
          background: var(--panel);
          border: 1px solid var(--line);
          border-radius: var(--radius);
          box-shadow: var(--shadow);
          padding: 16px;
        }
        .status-badge {
          display: inline-block;
          padding: 8px 12px;
          border-radius: 999px;
          background: {{ '#dff7ea' if status == 'Idle' else '#fff3cd' }};
          color: {{ '#1f7a4d' if status == 'Idle' else '#9a6700' }};
          font-weight: 700;
        }
        .button {
          display: inline-flex;
          align-items: center;
          justify-content: center;
          width: 100%;
          padding: 12px 14px;
          border-radius: 14px;
          text-decoration: none;
          font-weight: 700;
          border: none;
          cursor: pointer;
          background: var(--accent);
          color: white;
        }
        .button.secondary {
          background: #edf4fb;
          color: var(--accent);
        }
        .button:disabled {
          opacity: 0.6;
          cursor: not-allowed;
        }
        .section-title {
          margin: 18px 0 12px;
          font-size: 18px;
          font-weight: 800;
        }
        .logs {
          white-space: pre-wrap;
          font-family: Consolas, monospace;
          font-size: 12px;
          line-height: 1.4;
          min-height: 180px;
          max-height: 420px;
          overflow: auto;
          background: #0e2133;
          color: #d7e7f5;
          border-radius: 14px;
          padding: 14px;
        }
        .log-toolbar {
          display: flex;
          flex-wrap: wrap;
          gap: 10px;
          margin-bottom: 12px;
        }
        .inline-note {
          font-size: 12px;
          color: var(--muted);
        }
        .jobs {
          display: grid;
          gap: 16px;
        }
        .job {
          background: var(--panel);
          border: 1px solid var(--line);
          border-radius: 20px;
          box-shadow: var(--shadow);
          overflow: hidden;
        }
        .job-top {
          padding: 18px 18px 12px;
          border-bottom: 1px solid var(--line);
          background: linear-gradient(180deg, #fbfdff, #f3f7fb);
        }
        .job-title {
          font-size: 20px;
          font-weight: 800;
          margin: 0 0 8px;
        }
        .meta {
          display: flex;
          flex-wrap: wrap;
          gap: 8px;
        }
        .chip {
          background: #eaf2fa;
          color: #23486b;
          border-radius: 999px;
          padding: 6px 10px;
          font-size: 12px;
          font-weight: 700;
        }
        .job-body {
          display: grid;
          grid-template-columns: 1.2fr 0.8fr;
          gap: 16px;
          padding: 16px 18px 18px;
        }
        .stack {
          display: grid;
          gap: 12px;
        }
        .card {
          border: 1px solid var(--line);
          border-radius: 14px;
          padding: 14px;
          background: #fff;
        }
        .card h3 {
          margin: 0 0 10px;
          font-size: 14px;
          text-transform: uppercase;
          letter-spacing: 0.04em;
          color: var(--muted);
        }
        .links a {
          display: block;
          padding: 8px 0;
          color: var(--accent);
          text-decoration: none;
          word-break: break-word;
        }
        .reason {
          margin: 0;
          line-height: 1.55;
          color: var(--ink);
        }
        form {
          display: grid;
          gap: 10px;
        }
        select, textarea {
          width: 100%;
          border: 1px solid var(--line);
          border-radius: 12px;
          padding: 10px 12px;
          font: inherit;
          color: var(--ink);
          background: #fff;
        }
        textarea {
          min-height: 110px;
          resize: vertical;
        }
        .tiny {
          font-size: 12px;
          color: var(--muted);
        }
        @media (max-width: 900px) {
          .job-body {
            grid-template-columns: 1fr;
          }
        }
      </style>
    </head>
    <body>
      <div class="wrap">
        <div class="hero">
          <h1>Job Automation Dashboard</h1>
          <p>Run the search, review AI-evaluated jobs, download files, and track your application status from your phone.</p>
        </div>

        <div class="toolbar">
          <div class="panel">
            <div class="tiny">Runner</div>
            <div style="margin: 8px 0 14px;"><span class="status-badge">{{ status }}</span></div>
            <div class="tiny">Started: {{ run_state.started_at or "-" }}</div>
            <div class="tiny" style="margin-bottom: 10px;">Finished: {{ run_state.finished_at or "-" }}</div>
            <a class="button" href="/run">{{ "Already Running" if status == "Running" else "Run Job Search" }}</a>
          </div>
          <div class="panel">
            <div class="tiny">Jobs in dashboard</div>
            <div style="font-size: 30px; font-weight: 800; margin-top: 6px;">{{ jobs|length }}</div>
            <div class="tiny">Loaded from Supabase</div>
          </div>
          <div class="panel">
            <div class="tiny">Cover letters</div>
            <div style="font-size: 30px; font-weight: 800; margin-top: 6px;">{{ jobs|selectattr("has_cover_letter")|list|length }}</div>
            <div class="tiny">Jobs with generated letters</div>
          </div>
        </div>

        <div class="section-title">Recent Logs</div>
        <div class="panel">
          <div class="log-toolbar">
            <a class="button secondary" style="width:auto;" href="/logs" target="_blank" rel="noopener">Open Full Logs</a>
            <div class="inline-note">
              {% if status == "Running" %}
              Auto-refreshing every 5 seconds while running.
              {% else %}
              Refresh the page after a new run to see updated logs.
              {% endif %}
            </div>
          </div>
          <div class="logs">{{ recent_logs }}</div>
        </div>

        {% for section in grouped_sections %}
        <div class="section-title">{{ section.label }}</div>
        <div class="jobs">
          {% for job in section.jobs %}
          <div class="job">
            <div class="job-top">
              <div class="job-title">{{ job.title }}</div>
              <div class="meta">
                <span class="chip">{{ job.job_id or job.refnr }}</span>
                <span class="chip">{{ job.employer }}</span>
                <span class="chip">{{ job.city }}</span>
                <span class="chip">Date: {{ job.date }}</span>
                <span class="chip">Keyword score: {{ job.keyword_score }}</span>
                <span class="chip">AI match: {{ job.ai_match_score }}/10</span>
                <span class="chip">{{ job.application_status_label }}</span>
                <span class="chip">{{ "Cover letter ready" if job.has_cover_letter else "No cover letter" }}</span>
              </div>
            </div>
            <div class="job-body">
              <div class="stack">
                <div class="card links">
                  <h3>Links</h3>
                  {% if job.job_url %}
                  <a href="{{ job.job_url }}" target="_blank" rel="noopener">Open job listing</a>
                  {% else %}
                  <div class="tiny">No job URL recorded yet.</div>
                  {% endif %}

                  {% if job.has_cover_letter %}
                  <a href="{{ url_for('download_cover_letter', job_ai=(job.job_id or job.refnr), refnr=job.refnr) }}">Download cover letter</a>
                  {% else %}
                  <div class="tiny">No cover letter file.</div>
                  {% endif %}
                </div>

                <div class="card">
                  <h3>AI Reason</h3>
                  <p class="reason">{{ job.reason }}</p>
                </div>
              </div>

              <div class="stack">
                <div class="card">
                  <h3>Application Tracker</h3>
                  <form method="post" action="{{ url_for('update_job', refnr=job.refnr) }}">
                    <label class="tiny" for="status-{{ job.refnr }}">Application status</label>
                    <select id="status-{{ job.refnr }}" name="application_status">
                      <option value="not_applied" {{ "selected" if job.application_status == "not_applied" else "" }}>Not applied</option>
                      <option value="applied" {{ "selected" if job.application_status == "applied" else "" }}>Applied</option>
                      <option value="not_to_apply" {{ "selected" if job.application_status == "not_to_apply" else "" }}>Not to apply</option>
                    </select>

                    <label class="tiny" for="method-{{ job.refnr }}">Application method</label>
                    <input id="method-{{ job.refnr }}" name="application_method" value="{{ job.application_method }}" placeholder="Per E-Mail, Arbeitgeberseite, ..." />

                    <label class="tiny" for="result-{{ job.refnr }}">Application result</label>
                    <textarea id="result-{{ job.refnr }}" name="application_result" placeholder="Interview, rejection, no reply, accepted...">{{ job.application_result }}</textarea>

                    <label class="tiny" for="note-{{ job.refnr }}">Note</label>
                    <textarea id="note-{{ job.refnr }}" name="note" placeholder="Add your note here...">{{ job.note }}</textarea>

                    <button class="button secondary" type="submit">Save</button>
                    <div class="tiny">Job ID: {{ job.job_id or "Pending" }}</div>
                    <div class="tiny">Refnr: {{ job.refnr }}</div>
                    {% if job.updated_at %}
                    <div class="tiny">Last updated: {{ job.updated_at }}</div>
                    {% endif %}
                  </form>
                </div>
              </div>
            </div>
          </div>
          {% else %}
          <div class="panel">No jobs in this section yet.</div>
          {% endfor %}
        </div>
        {% endfor %}

        <div class="section-title">Application Results</div>
        <div class="jobs">
          {% for job in application_results %}
          <div class="job">
            <div class="job-top">
              <div class="job-title">{{ job.title }}</div>
              <div class="meta">
                <span class="chip">{{ job.job_id or job.refnr }}</span>
                <span class="chip">{{ job.employer }}</span>
                <span class="chip">{{ job.application_status_label }}</span>
              </div>
            </div>
            <div class="job-body">
              <div class="stack">
                <div class="card">
                  <h3>Application Outcome</h3>
                  <p class="reason">{{ job.application_result }}</p>
                </div>
              </div>
              <div class="stack">
                <div class="card">
                  <h3>Quick Links</h3>
                  <div class="links">
                    {% if job.job_url %}
                    <a href="{{ job.job_url }}" target="_blank" rel="noopener">Open job listing</a>
                    {% endif %}
                    {% if job.has_cover_letter %}
                    <a href="{{ url_for('download_cover_letter', job_ai=(job.job_id or job.refnr), refnr=job.refnr) }}">Download cover letter</a>
                    {% endif %}
                  </div>
                </div>
              </div>
            </div>
          </div>
          {% else %}
          <div class="panel">No application outcomes recorded yet.</div>
          {% endfor %}
        </div>
      </div>
    </body>
    </html>
    """
    return render_template_string(
        template,
        jobs=jobs,
        grouped_sections=grouped_sections,
        application_results=application_results,
        status=status,
        recent_logs=recent_logs,
        run_state=run_state,
    )


def filter_jobs_for_page(jobs, page_key):
    if page_key == "all":
        return jobs
    if page_key == "results":
        return [job for job in jobs if (job.get("application_result") or "").strip()]
    if page_key == "applied":
        return [job for job in jobs if (job.get("application_status") or "") == "applied"]
    if page_key == "not_applied":
        return [job for job in jobs if (job.get("application_status") or "") == "not_applied"]
    return jobs


def render_jobs_page(page_title, page_key):
    jobs = filter_jobs_for_page(load_jobs(), page_key)
    running = is_running()
    run_state = load_run_state()
    status = "Läuft" if running else run_state.get("status", "idle").title()
    return render_template_string(
        """
        <!doctype html>
        <html lang="de">
        <head>
          <meta charset="utf-8">
          <meta name="viewport" content="width=device-width, initial-scale=1">
          <title>{{ page_title }}</title>
          <style>
            :root {
              --bg-0:#05070d;
              --bg-1:#0a1021;
              --bg-2:#111a34;
              --panel:#111b35;
              --panel-2:#0d1530;
              --line:#2b3c64;
              --ink:#edf3ff;
              --muted:#9cb0d8;
              --primary:#56a4ff;
              --primary-2:#7a6dff;
              --ok:#4ad192;
              --warn:#ffc863;
              --radius:18px;
              --shadow:0 18px 42px rgba(4,7,16,.5);
            }
            *{box-sizing:border-box}
            body{
              margin:0;
              color:var(--ink);
              font-family:"Aptos","Segoe UI",Arial,sans-serif;
              background:
                radial-gradient(1100px 600px at 100% -8%, rgba(122,109,255,.22), transparent 60%),
                radial-gradient(900px 520px at -8% 24%, rgba(86,164,255,.18), transparent 58%),
                linear-gradient(180deg,var(--bg-0),var(--bg-1) 35%,var(--bg-2));
            }
            .shell{max-width:1360px;margin:0 auto;padding:18px 16px 34px}
            .topbar{
              position:sticky;top:0;z-index:40;background:rgba(5,7,13,.86);backdrop-filter:blur(10px);
              border-bottom:1px solid rgba(156,176,216,.2);margin:-18px -16px 14px;padding:14px 16px 12px;
            }
            .nav{display:flex;gap:8px;flex-wrap:wrap}
            .nav a{
              padding:10px 14px;border-radius:11px;border:1px solid rgba(156,176,216,.3);text-decoration:none;
              color:var(--ink);background:rgba(17,27,53,.86);font-weight:700;font-size:13px;letter-spacing:.02em;
            }
            .nav a.active{background:linear-gradient(135deg,var(--primary),var(--primary-2));border-color:transparent;color:#fff}
            .hero{
              display:grid;grid-template-columns:1fr auto;gap:12px;align-items:center;
              background:linear-gradient(145deg,#15264d,#0f1733 64%);
              border:1px solid rgba(156,176,216,.28);
              border-radius:22px;
              padding:18px;
              box-shadow:var(--shadow);
            }
            .hero h1{margin:0 0 6px;font-size:30px;letter-spacing:-.02em}
            .hero p{margin:0;color:var(--muted);font-size:14px}
            .status{padding:8px 12px;border-radius:999px;font-size:12px;font-weight:800;letter-spacing:.04em}
            .status.ok{background:rgba(74,209,146,.16);color:var(--ok);border:1px solid rgba(74,209,146,.38)}
            .status.warn{background:rgba(255,200,99,.16);color:var(--warn);border:1px solid rgba(255,200,99,.38)}
            .stats{margin-top:12px;display:grid;grid-template-columns:repeat(4,minmax(0,1fr));gap:10px}
            .metric{background:rgba(17,27,53,.75);border:1px solid rgba(156,176,216,.26);border-radius:14px;padding:12px}
            .metric .k{font-size:11px;color:var(--muted);text-transform:uppercase;letter-spacing:.06em}
            .metric .v{font-size:22px;font-weight:800;margin-top:6px}
            .cards{display:grid;gap:12px;margin-top:14px}
            details.card{
              border:1px solid rgba(156,176,216,.26);border-radius:18px;overflow:hidden;
              background:linear-gradient(180deg,var(--panel),var(--panel-2));box-shadow:var(--shadow);
            }
            details > summary{list-style:none;cursor:pointer;padding:14px 16px}
            details > summary::-webkit-details-marker{display:none}
            .sum{display:grid;grid-template-columns:1fr auto;gap:10px;align-items:center}
            .ttl{font-size:18px;font-weight:800;line-height:1.3}
            .sub{color:var(--muted);font-size:12px;margin-top:4px}
            .chips{display:flex;gap:6px;flex-wrap:wrap;justify-content:flex-end}
            .chip{border:1px solid rgba(156,176,216,.32);background:rgba(9,13,27,.38);border-radius:999px;padding:4px 8px;font-size:11px;font-weight:700;color:#d8e4ff}
            .body{border-top:1px solid rgba(156,176,216,.24);padding:14px 16px 16px;display:grid;grid-template-columns:1.12fr .88fr;gap:12px}
            .panel{border:1px solid rgba(156,176,216,.25);background:rgba(9,13,27,.32);border-radius:14px;padding:12px}
            .panel h3{margin:0 0 10px;font-size:12px;text-transform:uppercase;letter-spacing:.08em;color:#b5c7ea}
            .txt{color:#c8d6f4;font-size:13px;line-height:1.55}
            .actions{display:flex;gap:8px;flex-wrap:wrap;margin-top:10px}
            .btn{display:inline-flex;align-items:center;justify-content:center;padding:10px 12px;border-radius:10px;border:1px solid transparent;background:linear-gradient(135deg,var(--primary),var(--primary-2));color:#fff;text-decoration:none;font-size:12px;font-weight:800;letter-spacing:.03em;cursor:pointer}
            .btn.secondary{background:rgba(17,27,53,.82);border-color:rgba(156,176,216,.36);color:#dbe6ff}
            form{display:grid;gap:8px}
            label{font-size:12px;font-weight:700;color:#b8c8ea}
            select,input,textarea{width:100%;border-radius:10px;border:1px solid rgba(156,176,216,.38);padding:10px 11px;background:rgba(8,12,25,.5);color:var(--ink);font:inherit}
            textarea{min-height:86px;resize:vertical}
            .empty{border:1px dashed rgba(156,176,216,.36);background:rgba(17,27,53,.5);border-radius:14px;padding:30px;text-align:center;color:var(--muted)}
            @media (max-width:1060px){
              .stats{grid-template-columns:repeat(2,minmax(0,1fr))}
              .body{grid-template-columns:1fr}
              .sum{grid-template-columns:1fr}
              .chips{justify-content:flex-start}
              .hero{grid-template-columns:1fr}
              .hero h1{font-size:25px}
            }
          </style>
        </head>
        <body>
          <div class="shell">
            <header class="topbar">
              <nav class="nav" aria-label="Hauptnavigation">
                <a href="{{ url_for('search_page') }}">Suche starten</a>
                <a href="{{ url_for('applied_jobs_page') }}" class="{{ 'active' if page_key=='applied' else '' }}">Beworben</a>
                <a href="{{ url_for('not_applied_jobs_page') }}" class="{{ 'active' if page_key=='not_applied' else '' }}">Nicht beworben</a>
                <a href="{{ url_for('results_jobs_page') }}" class="{{ 'active' if page_key=='results' else '' }}">Ergebnisse</a>
                <a href="{{ url_for('all_jobs_page') }}" class="{{ 'active' if page_key=='all' else '' }}">Alle Stellen</a>
                <a href="{{ url_for('settings_page') }}">Einstellungen</a>
              </nav>
            </header>

            <section class="hero">
              <div>
                <h1>{{ page_title }}</h1>
                <p>Vollständig neues Dashboard: klare Hierarchie, schnelle Übersicht und fokussierte Detailbearbeitung pro Stelle.</p>
              </div>
              <span class="status {{ 'warn' if status == 'Läuft' else 'ok' }}">{{ status }}</span>
            </section>

            <section class="stats">
              <article class="metric"><div class="k">Anzeigen</div><div class="v">{{ jobs|length }}</div></article>
              <article class="metric"><div class="k">Mit Anschreiben</div><div class="v">{{ jobs|selectattr('has_cover_letter')|list|length }}</div></article>
              <article class="metric"><div class="k">Status</div><div class="v">{{ status }}</div></article>
              <article class="metric"><div class="k">Aktualisiert</div><div class="v" style="font-size:14px;line-height:1.3;">{{ run_state.finished_at or "-" }}</div></article>
            </section>

            <section class="cards">
              {% for job in jobs %}
              <details class="card">
                <summary>
                  <div class="sum">
                    <div>
                      <div class="ttl">{{ job.title or "Unbenannte Stelle" }}</div>
                      <div class="sub">{{ job.employer }} | {{ job.city }} | {{ job.date or "Kein Datum" }}</div>
                    </div>
                    <div class="chips">
                      <span class="chip">{{ job.job_id or job.refnr }}</span>
                      <span class="chip">AI {{ job.ai_match_score or "-" }}/10</span>
                      <span class="chip">Score {{ job.keyword_score or "-" }}</span>
                      <span class="chip">{{ job.application_status_label }}</span>
                    </div>
                  </div>
                </summary>
                <div class="body">
                  <section class="panel">
                    <h3>Stelleninfos</h3>
                    <div class="sub"><strong>Refnr:</strong> {{ job.refnr }}</div>
                    <p class="txt" style="margin:10px 0 0;">{{ job.reason or "Noch keine AI-Begründung gespeichert." }}</p>
                    <div class="actions">
                      {% if job.job_url %}<a class="btn" href="{{ job.job_url }}" target="_blank" rel="noopener">Stellenquelle öffnen</a>{% endif %}
                      {% if job.has_cover_letter %}<a class="btn secondary" href="{{ url_for('download_cover_letter', job_ai=(job.job_id or job.refnr), refnr=job.refnr) }}">Anschreiben herunterladen</a>{% endif %}
                    </div>
                  </section>
                  <section class="panel">
                    <h3>Bewerbung verwalten</h3>
                    <form method="post" action="{{ url_for('update_job', refnr=job.refnr) }}">
                      <label>Status</label>
                      <select name="application_status">
                        <option value="not_applied" {{ "selected" if job.application_status == "not_applied" else "" }}>Nicht beworben</option>
                        <option value="applied" {{ "selected" if job.application_status == "applied" else "" }}>Beworben</option>
                        <option value="not_to_apply" {{ "selected" if job.application_status == "not_to_apply" else "" }}>Nicht bewerben</option>
                      </select>
                      <label>Bewerbungsweg</label>
                      <input name="application_method" value="{{ job.application_method }}" placeholder="z.B. Per E-Mail" />
                      <label>Ergebnis</label>
                      <textarea name="application_result" placeholder="Ergebnis / Rückmeldung">{{ job.application_result }}</textarea>
                      <label>Notiz</label>
                      <textarea name="note" placeholder="Interne Notiz">{{ job.note }}</textarea>
                      <button class="btn" type="submit">Änderungen speichern</button>
                    </form>
                  </section>
                </div>
              </details>
              {% else %}
              <div class="empty">In diesem Bereich sind aktuell keine Stellen vorhanden.</div>
              {% endfor %}
            </section>
          </div>
        </body>
        </html>
        """,
        page_title=page_title,
        page_key=page_key,
        jobs=jobs,
        status=status,
        run_state=run_state,
    )

@app.get("/search")
def search_page():
    running = is_running()
    run_state = load_run_state()
    status = "Läuft" if running else run_state.get("status", "idle").title()
    logs = LOG_FILE.read_text(encoding="utf-8", errors="replace") if LOG_FILE.exists() else "Noch keine Logs vorhanden."
    recent_logs = "\n".join(clean_log_content(logs).splitlines()[-120:])
    return render_template_string(
        """
        <!doctype html>
        <html lang="de"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1"><title>Suche starten</title>
        {% if status == "Läuft" %}<meta http-equiv="refresh" content="5">{% endif %}
        <style>
          :root{
            --bg:#090d1b;--bg2:#0f1630;--line:#273456;--ink:#f2f6ff;--muted:#9fb0d4;
            --primary:#55a0ff;--primary2:#6d6dff;--ok:#44d18f;--warn:#ffc861;--shadow:0 16px 38px rgba(5,8,20,.45);
          }
          body{
            margin:0;font-family:"Aptos","Segoe UI",Arial,sans-serif;color:var(--ink);
            background:
              radial-gradient(1200px 600px at 90% -10%, rgba(109,109,255,.20), transparent 55%),
              radial-gradient(900px 520px at -10% 20%, rgba(85,160,255,.18), transparent 55%),
              linear-gradient(180deg,var(--bg),var(--bg2));
          }
          .shell{max-width:1160px;margin:0 auto;padding:18px 16px 34px}
          .topbar{position:sticky;top:0;z-index:40;background:rgba(5,7,13,.86);backdrop-filter:blur(10px);border-bottom:1px solid rgba(156,176,216,.2);margin:-18px -16px 14px;padding:14px 16px 12px}
          .nav{display:flex;gap:8px;flex-wrap:wrap}
          .nav a{padding:10px 14px;border-radius:11px;border:1px solid rgba(156,176,216,.3);text-decoration:none;color:var(--ink);background:rgba(17,27,53,.86);font-weight:700;font-size:13px;letter-spacing:.02em}
          .nav a.active{background:linear-gradient(135deg,var(--primary),var(--primary2));border-color:transparent;color:#fff}
          .hero{display:grid;grid-template-columns:1fr auto;gap:12px;align-items:center;background:linear-gradient(145deg,#15264d,#0f1733 64%);border:1px solid rgba(156,176,216,.28);border-radius:22px;padding:18px;box-shadow:var(--shadow)}
          .hero h1{margin:0 0 6px;font-size:30px}
          .hero p{margin:0;color:var(--muted)}
          .status{padding:8px 12px;border-radius:999px;font-size:12px;font-weight:800;letter-spacing:.04em}
          .status.ok{background:rgba(74,209,146,.16);color:var(--ok);border:1px solid rgba(74,209,146,.38)}
          .status.warn{background:rgba(255,200,99,.16);color:var(--warn);border:1px solid rgba(255,200,99,.38)}
          .actions{display:flex;gap:8px;flex-wrap:wrap;margin-top:12px}
          .btn{display:inline-flex;align-items:center;justify-content:center;padding:11px 14px;border-radius:10px;border:1px solid transparent;background:linear-gradient(135deg,var(--primary),var(--primary2));color:#fff;text-decoration:none;font-size:12px;font-weight:800;letter-spacing:.03em}
          .meta{margin-top:8px;color:var(--muted);font-size:13px}
          .panel{margin-top:12px;background:linear-gradient(180deg,#101a35,#0d1530);border:1px solid rgba(156,176,216,.26);border-radius:18px;box-shadow:var(--shadow);padding:14px}
          .panel h2{margin:0 0 10px;font-size:16px}
          .logs{white-space:pre-wrap;background:rgba(9,13,27,.5);border:1px solid rgba(156,176,216,.26);color:#d7e7f5;border-radius:12px;padding:13px;font-family:Consolas,monospace;font-size:12px;min-height:340px;line-height:1.42}
          @media (max-width:980px){.hero{grid-template-columns:1fr}.hero h1{font-size:24px}}
        </style></head><body><div class="shell">
          <header class="topbar">
            <nav class="nav">
              <a href="{{ url_for('search_page') }}" class="active">Suche starten</a>
              <a href="{{ url_for('applied_jobs_page') }}">Beworben</a>
              <a href="{{ url_for('not_applied_jobs_page') }}">Nicht beworben</a>
              <a href="{{ url_for('results_jobs_page') }}">Ergebnisse</a>
              <a href="{{ url_for('all_jobs_page') }}">Alle Stellen</a>
              <a href="{{ url_for('settings_page') }}">Einstellungen</a>
            </nav>
          </header>
          <section class="hero">
            <div>
              <h1>Suche starten</h1>
              <p>Führe den kompletten Lauf aus und verfolge den Fortschritt live im Logfenster.</p>
              <div class="meta">Gestartet: {{ run_state.started_at or "-" }} | Beendet: {{ run_state.finished_at or "-" }}</div>
              <div class="actions"><a class="btn" href="/run">{{ "Läuft bereits" if status == "Läuft" else "Jobsuche jetzt starten" }}</a></div>
            </div>
            <span class="status {{ 'warn' if status == 'Läuft' else 'ok' }}">{{ status }}</span>
          </section>
          <section class="panel">
            <h2>Live-Logs</h2>
            <div class="logs">{{ recent_logs }}</div>
          </section>
        </div></body></html>
        """,
        status=status,
        run_state=run_state,
        recent_logs=recent_logs,
    )

@app.get("/jobs/applied")
def applied_jobs_page():
    return render_jobs_page("Applied Jobs", "applied")


@app.get("/jobs/not-applied")
def not_applied_jobs_page():
    return render_jobs_page("Not Applied Jobs", "not_applied")


@app.get("/jobs/results")
def results_jobs_page():
    return render_jobs_page("Results", "results")


@app.get("/jobs/all")
def all_jobs_page():
    return render_jobs_page("All Jobs", "all")


@app.get("/settings")
def settings_page():
    terms = "\n".join(load_search_terms())
    return render_template_string(
        """
        <!doctype html>
        <html lang="de"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1"><title>Einstellungen</title>
        <style>
          :root{
            --bg:#090d1b;--bg2:#0f1630;--line:#273456;--ink:#f2f6ff;--muted:#9fb0d4;
            --primary:#55a0ff;--primary2:#6d6dff;--shadow:0 16px 38px rgba(5,8,20,.45);
          }
          body{
            margin:0;font-family:"Aptos","Segoe UI",Arial,sans-serif;color:var(--ink);
            background:
              radial-gradient(1200px 600px at 90% -10%, rgba(109,109,255,.20), transparent 55%),
              radial-gradient(900px 520px at -10% 20%, rgba(85,160,255,.18), transparent 55%),
              linear-gradient(180deg,var(--bg),var(--bg2));
          }
          .shell{max-width:1000px;margin:0 auto;padding:18px 16px 34px}
          .topbar{position:sticky;top:0;z-index:40;background:rgba(5,7,13,.86);backdrop-filter:blur(10px);border-bottom:1px solid rgba(156,176,216,.2);margin:-18px -16px 14px;padding:14px 16px 12px}
          .nav{display:flex;gap:8px;flex-wrap:wrap}
          .nav a{padding:10px 14px;border-radius:11px;border:1px solid rgba(156,176,216,.3);text-decoration:none;color:var(--ink);background:rgba(17,27,53,.86);font-weight:700;font-size:13px;letter-spacing:.02em}
          .nav a.active{background:linear-gradient(135deg,var(--primary),var(--primary2));border-color:transparent;color:#fff}
          .hero{background:linear-gradient(145deg,#15264d,#0f1733 64%);border:1px solid rgba(156,176,216,.28);border-radius:22px;padding:18px;box-shadow:var(--shadow)}
          .hero h1{margin:0 0 6px;font-size:30px}
          .hero p{margin:0;color:var(--muted)}
          .panel{margin-top:12px;background:linear-gradient(180deg,#101a35,#0d1530);border:1px solid rgba(156,176,216,.26);border-radius:18px;box-shadow:var(--shadow);padding:14px}
          label{display:block;font-size:12px;color:#b8c7e9;margin-bottom:6px;font-weight:800;letter-spacing:.05em}
          textarea{width:100%;min-height:280px;background:rgba(9,13,27,.45);border:1px solid rgba(156,176,216,.36);border-radius:12px;padding:12px;color:var(--ink);font:inherit}
          .actions{margin-top:10px;display:flex;gap:8px;flex-wrap:wrap}
          .btn{display:inline-flex;align-items:center;justify-content:center;padding:10px 13px;border-radius:10px;text-decoration:none;background:linear-gradient(135deg,var(--primary),var(--primary2));color:#fff;font-weight:800;font-size:12px;letter-spacing:.03em;border:1px solid transparent;cursor:pointer}
          .btn.secondary{background:rgba(16,26,53,.8);border-color:rgba(156,176,216,.35);color:#d9e6ff}
        </style></head><body><div class="shell">
          <header class="topbar">
            <nav class="nav">
              <a href="{{ url_for('search_page') }}">Suche starten</a>
              <a href="{{ url_for('applied_jobs_page') }}">Beworben</a>
              <a href="{{ url_for('not_applied_jobs_page') }}">Nicht beworben</a>
              <a href="{{ url_for('results_jobs_page') }}">Ergebnisse</a>
              <a href="{{ url_for('all_jobs_page') }}">Alle Stellen</a>
              <a href="{{ url_for('settings_page') }}" class="active">Einstellungen</a>
            </nav>
          </header>
          <section class="hero">
            <h1>Einstellungen</h1>
            <p>Bearbeite Suchbegriffe und exportiere die Bewerbungsübersicht aus einem zentralen Bereich.</p>
          </section>
          <section class="panel">
            <form method="post" action="{{ url_for('save_settings_page') }}">
              <label>SUCHBEGRIFFE (EIN BEGRIFF PRO ZEILE)</label>
              <textarea name="search_terms_text">{{ terms }}</textarea>
              <div class="actions">
                <button class="btn" type="submit">Suchbegriffe speichern</button>
                <a class="btn secondary" href="{{ url_for('download_applications_summary') }}">Bewerbungsübersicht herunterladen</a>
              </div>
            </form>
          </section>
        </div></body></html>
        """,
        terms=terms,
    )

@app.post("/settings")
def save_settings_page():
    text = request.form.get("search_terms_text", "")
    save_search_terms_text(text)
    return redirect(url_for("settings_page"))


@app.get("/run")
def run_script():
    global current_process

    if is_running():
        return redirect(url_for("search_page"))

    start_search_process()
    return redirect(url_for("search_page"))


def start_search_process():
    global current_process

    log_handle = open(LOG_FILE, "w", encoding="utf-8")
    log_handle.write("Starting job search...\n")
    log_handle.flush()
    env = os.environ.copy()
    env["PYTHONUNBUFFERED"] = "1"
    try:
        current_process = subprocess.Popen(
            [sys.executable, "-u", str(BASE_DIR / "main.py")],
            cwd=BASE_DIR,
            stdout=log_handle,
            stderr=subprocess.STDOUT,
            text=True,
            bufsize=1,
            env=env,
        )
        save_run_state("running")
        return True, None
    except Exception as exc:
        log_handle.write(f"Failed to start job search: {exc}\n")
        log_handle.close()
        save_run_state("failed", returncode=-1)
        current_process = None
        return False, str(exc)


@app.post("/api/run")
@app.get("/api/run")
def api_run_script():
    if is_running():
        return jsonify({"started": False, "status": "already_running"})

    started, error = start_search_process()
    if not started:
        return jsonify({"started": False, "status": "failed", "error": error}), 500
    return jsonify({"started": True, "status": "running"})


@app.get("/logs")
def logs():
    running = is_running()
    run_state = load_run_state()
    raw_content = LOG_FILE.read_text(encoding="utf-8", errors="replace") if LOG_FILE.exists() else "No log file yet."
    content = clean_log_content(raw_content)
    return render_template_string(
        """
        <!doctype html>
        <html lang="en">
        <head>
          <meta charset="utf-8">
          <meta name="viewport" content="width=device-width, initial-scale=1">
          <title>Run Logs</title>
          {% if running %}
          <meta http-equiv="refresh" content="5">
          {% endif %}
          <style>
            body {
              margin: 0;
              padding: 18px;
              background: #0c1b2a;
              color: #d7e7f5;
              font-family: Consolas, monospace;
            }
            .top {
              margin-bottom: 14px;
              font-family: "Segoe UI", Arial, sans-serif;
            }
            a {
              color: #8dc5ff;
            }
            pre {
              white-space: pre-wrap;
              line-height: 1.45;
              font-size: 13px;
            }
          </style>
        </head>
        <body>
          <div class="top">
            <div><strong>Status:</strong> {{ run_state.status }}</div>
            <div><strong>Started:</strong> {{ run_state.started_at or "-" }}</div>
            <div><strong>Finished:</strong> {{ run_state.finished_at or "-" }}</div>
            <div><a href="/">Back to dashboard</a></div>
          </div>
          <pre>{{ content }}</pre>
        </body>
        </html>
        """,
        content=content,
        running=running,
        run_state=run_state,
    )


@app.get("/api/logs")
def api_logs():
    raw_content = LOG_FILE.read_text(encoding="utf-8", errors="replace") if LOG_FILE.exists() else "No log file yet."
    content = clean_log_content(raw_content)
    return app.response_class(content, mimetype="text/plain; charset=utf-8")


@app.get("/api/search-terms")
def get_search_terms():
    return jsonify({"terms": load_search_terms(), "text": "\n".join(load_search_terms())})


@app.post("/api/search-terms")
def update_search_terms():
    payload = request.get_json(silent=True) or {}
    text = str(payload.get("text", "")).strip()
    if not text and isinstance(payload.get("terms"), list):
        text = "\n".join(str(term) for term in payload["terms"])
    terms = save_search_terms_text(text)
    return jsonify({"saved": True, "terms": terms, "text": "\n".join(terms)})


@app.get("/download/applications-summary.xlsx")
def download_applications_summary():
    jobs = [
        job
        for job in load_jobs()
        if job.get("application_status") == "applied"
    ]
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Bewerbungen"
    sheet.append(
        [
            "Name der Firma",
            "Straße",
            "Postleitzahl",
            "Ort",
            "Als was hab ich mich beworben?",
            "Wann",
            "Wie",
            "Ergebnis",
        ]
    )
    for job in jobs:
        sheet.append(
            [
                job.get("employer", ""),
                job.get("employer_street", ""),
                job.get("employer_postal_code", ""),
                job.get("employer_city") or job.get("city", ""),
                job.get("title", ""),
                date_only(job.get("date")),
                job.get("application_method", ""),
                job.get("application_result", ""),
            ]
        )

    for column in sheet.columns:
        max_length = max(len(str(cell.value or "")) for cell in column)
        sheet.column_dimensions[column[0].column_letter].width = min(max(max_length + 2, 12), 42)

    buffer = BytesIO()
    workbook.save(buffer)
    buffer.seek(0)
    return send_file(
        buffer,
        as_attachment=True,
        download_name="bewerbungen-zusammenfassung.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.post("/jobs/<refnr>/update")
def update_job(refnr):
    client = get_supabase_client()
    if client is not None:
        application_status = request.form.get("application_status", "not_applied")
        if application_status == "to_apply":
            application_status = "not_applied"
        existing = fetch_job(refnr) or {}
        applied_at = existing.get("applied_at")
        if application_status == "applied" and not applied_at:
            applied_at = datetime.now().isoformat()
        client.table(SUPABASE_TABLE).update(
            {
                "application_status": application_status,
                "application_method": request.form.get("application_method", "").strip(),
                "application_result": request.form.get("application_result", "").strip(),
                "applied_at": applied_at,
                "note": request.form.get("note", "").strip(),
                "created_at": existing.get("created_at"),
                "updated_at": datetime.now().isoformat(),
            }
        ).eq("refnr", str(refnr)).execute()
    referer = request.headers.get("Referer")
    if referer:
        return redirect(referer)
    return redirect(url_for("all_jobs_page"))


@app.patch("/api/jobs/<refnr>")
def update_job_api(refnr):
    client = get_supabase_client()
    if client is None:
        return jsonify({"error": "Supabase is not configured"}), 500

    payload = request.get_json(silent=True) or {}
    application_status = payload.get("application_status", "not_applied")
    if application_status == "to_apply":
        application_status = "not_applied"
    application_method = str(payload.get("application_method", "")).strip()
    application_result = str(payload.get("application_result", "")).strip()
    note = str(payload.get("note", "")).strip()
    updated_at = payload.get("updated_at") or datetime.now().isoformat()
    existing = fetch_job(refnr) or {}
    applied_at = payload.get("applied_at") or existing.get("applied_at")
    if application_status == "applied" and not applied_at:
        applied_at = updated_at

    response = client.table(SUPABASE_TABLE).update(
        {
            "application_status": application_status,
            "application_method": application_method,
            "application_result": application_result,
            "applied_at": applied_at,
            "note": note,
            "created_at": existing.get("created_at"),
            "updated_at": updated_at,
        }
    ).eq("refnr", str(refnr)).execute()
    return jsonify({"updated": True, "data": response.data or []})


@app.get("/download/cover-letter/<path:job_ai>")
def download_cover_letter(job_ai):
    refnr = request.args.get("refnr", "").strip()
    row = fetch_job(refnr) if refnr else None
    if row is None:
        row = fetch_job_by_job_id(job_ai)
    if not row:
        return "Cover letter not found", 404

    cover_letter_text = (row.get("cover_letter_text") or "").strip()
    if not cover_letter_text:
        return "Cover letter text not found", 404
    if not TEMPLATE_PATH.exists():
        return "Cover letter template not found", 500

    download_id = str(job_ai or row.get("job_id") or row.get("refnr") or "job").strip()
    file_name = f"hadi-komail-anschreibe-{download_id}.docx"
    safe_name = "".join(ch if ch.isalnum() or ch in ("-", "_", ".") else "_" for ch in file_name)

    document = Document(TEMPLATE_PATH)
    fill_cover_letter_template(document, row)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return send_file(
        buffer,
        as_attachment=True,
        download_name=safe_name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", "5000")), debug=False)


